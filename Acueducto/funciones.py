# FUNCION GENERAR TOKEN
from datetime import datetime, timedelta
from jose import jwt
from fpdf import FPDF
from fastapi import (
    FastAPI,
    Request,
    Form,
    status,
    Depends,
    HTTPException,
    Cookie,
    Query,
)
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from cruds.ReunionesCrud import obtenerReuAdmin
from fastapi.responses import JSONResponse, RedirectResponse
from models import Token, Usuario, Empresa, Vivienda, Reunion
from sqlalchemy.orm import Session
from docx2pdf import convert
import PyPDF2
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement

app = FastAPI()

# Agregando los archivos estaticos que están en la carpeta dist del proyecto
app.mount("/static", StaticFiles(directory="public/dist"), name="static")

template = Jinja2Templates(directory="public/templates")

SECRET_KEY = "sd45g4f45SWFGVHHuoyiad4F5SFD65V4SFDVOJWNHACUfwghdfvcguDCwfghezxhAzAKHGFBJYTFdkjfghtjkdgb"


def generar_token(usuario_id: str):
    expiration = datetime.utcnow() + timedelta(hours=20)  # Token expira en 20 hora
    payload = {"sub": usuario_id, "exp": expiration}
    token = jwt.encode(payload, SECRET_KEY, algorithm="HS256")
    return token

# FUNCION VERIFICAR TOKEN


def verificar_token(token: str, db):
    if token is None:
        return None
    consultar = db.query(Token).filter(Token.token == token).first()
    if consultar:
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms="HS256")
            return payload.get("sub")
        except jwt.ExpiredSignatureError:  # Token ha expirado
            deleteToken = db.query(Token).filter(Token.token == token).first()
            if deleteToken:
                db.delete(deleteToken)
                db.commit()
            return None
        except jwt.JWTError:  # Token inválido
            return None
    else:  # Token no válido
        return None


# FUNCION PARA OBTENER EL ROL DE USUARIO
def get_rol(id_usuario, db):
    if id_usuario:
        usuario = db.query(Usuario).filter(
            Usuario.id_usuario == id_usuario).first()
        if usuario:
            return usuario.rol
        else:
            return None
    else:
        return None

# FUNCION PARA OBTENER LA EMPRESA DEL USUARIO


def get_empresa(id_usuario, db):
    if id_usuario:
        usuario = db.query(Usuario).filter(
            Usuario.id_usuario == id_usuario).first()
        if usuario:
            return usuario.empresa
        else:
            return None
    else:
        return None


# FUNCION PARA OBTENER LOS DATOS DE USUARIO
def get_datos_usuario(id_usuario, db):
    if id_usuario:
        usuario = db.query(Usuario).filter(
            Usuario.id_usuario == id_usuario).first()
        if usuario:
            datos_usuario = {
                "id_usuario": usuario.id_usuario,
                "nom_usuario": usuario.nom_usuario,
                "apellido_usuario": usuario.apellido_usuario,
                "num_doc": usuario.num_doc,
                "direccion": usuario.direccion,
                "municipio": usuario.municipio,
                "rol": usuario.rol,
                "estado": usuario.estado,
                "empresa": usuario.empresa,
                "correo": usuario.correo
                # Agrega otros campos del usuario
            }
            return datos_usuario
        else:
            return None
    else:
        return None

# CAMBIAR LOS CAMPOS "[]" POR VALORES DE FORMULARIO


def reemplazar_texto(docx_path, datos):
    document = Document(docx_path)

    for paragraph in document.paragraphs:
        for campo, valor in datos.items():
            if campo in paragraph.text:
                paragraph.text = paragraph.text.replace(campo, valor)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for campo, valor in datos.items():
                        if campo in paragraph.text:
                            paragraph.text = paragraph.text.replace(campo, valor)

    return document

# CONVIERTE EL DOCUMENTO DOCX A PDF


def convertir_a_pdf(docx_path, pdf_path):
    try:
        # Llama a la función convert para convertir el documento DOCX a PDF
        convert(docx_path, pdf_path)
        return True  # La conversión fue exitosa
    except Exception as e:
        print(f"Error al convertir a PDF: {str(e)}")
        return False  # Ocurrió un error durante la conversión


# OBTENER DATOS EMPRESA:
def get_datos_empresa(id_empresa, db):
    if id_empresa:
        empresa = db.query(Empresa).filter(
            Empresa.id_empresa == id_empresa).first()
        if empresa:
            datos_empresa = {
                "id_empresa": empresa.id_empresa,
                "nom_empresa": empresa.nom_empresa,
                "direccion_empresa": empresa.direccion_empresa,
                "tel_fijo": empresa.tel_fijo,
                "tel_cel": empresa.tel_cel,
                "email": empresa.email,
                "estado": empresa.estado,
            }
            return datos_empresa
        else:
            return None
    else:
        return None


# OBTERNER DATOS EMPRESA (TODAS LAS REGISTRADAS):

def get_datos_empresas(db) -> list[str]:
    nom_empresas = db.query(Empresa.nom_empresa).all()
    return [nombre[0] for nombre in nom_empresas]


# OBTENER DATOS REUNIONES:
def get_datos_reuniones(id_reunion, db):
    if id_reunion:
        reunion = db.query(Reunion).filter(
            Reunion.id_reunion == id_reunion).first()
        if reunion:
            datos_reunion = {
                "id_reunion": reunion.id_reunion,
                "nom_reunion": reunion.nom_reunion,
                "id_empresa": reunion.id_empresa,
                "fecha": reunion.fecha,
                "url_asistencia": reunion.url_asistencia,
            }
            return datos_reunion
        else:
            return None
    else:
        return None

# FUNCION PARA ELIMINAR EL CACHE (HEADERS) 4/10/2023

def elimimar_cache():
    headers = {
        "Cache-Control": "no-store, must-revalidate",
        "Pragma": "no-cache",
    }
    return headers


def get_datos_vivienda(id_vivienda, db):
    if id_vivienda:
        vivienda = db.query(Vivienda).filter(
            Vivienda.id_inmueble == id_vivienda).first()
        if vivienda:
            datos_vivienda = {
                "id_vivienda": vivienda.id_inmueble,
                "id_usuario": vivienda.id_usuario,
                "direccion": vivienda.direccion,
                "estrato": vivienda.estrato,
                "uso": vivienda.uso,
                "numero_residentes": vivienda.numero_residentes,
            }
            return datos_vivienda
        else:
            return None
    else:
        return None


def get_viviendas(id_usuario, db):
    if id_usuario:
        viviendas = db.query(Vivienda).filter(
            Vivienda.id_usuario == id_usuario).all()
        if viviendas:
            return viviendas
        else:
            return None
    else:
        return None
    
def get_viviendas_empresa(id_empresa, db):
    if id_empresa:
        viviendas = db.query(Vivienda).join(Usuario).filter(
            Usuario.empresa == id_empresa).all()
        if viviendas:
            return viviendas
        else:
            return None
    else:
        return None
    
    # comentario

#CALCULAR CUORUM

def calcularCuorum(
    db:Session,
    token_valido:str,
    request:Request,
    cantidadAsistentes: int,
    reunion_1:str,
):  
    
    if token_valido:
        
        usuario = (
                db.query(Usuario).filter(
                    Usuario.id_usuario == token_valido).first()
            )
        
        query_subs = db.query(Usuario).filter(Usuario.rol == "Suscriptor", Usuario.empresa == usuario.empresa).count()

        if cantidadAsistentes > query_subs/2:
            
            sacarCuorum = True
        else:
            sacarCuorum = False
            
        update_reunion = (
            db.query(Reunion).filter_by(id_reunion=reunion_1).first()
        )
        if sacarCuorum:
            if update_reunion:
                update_reunion.cuorum = 1
                db.commit()
        else:
            if update_reunion:
                update_reunion.cuorum = 0
                db.commit()

        return sacarCuorum

    else:
        return RedirectResponse(url="/", status_code=status.HTTP_303_SEE_OTHER)


